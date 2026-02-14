
from __future__ import annotations
import re, io, time, os
from functools import wraps
from typing import Dict, List, Any, Optional
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import pandas as pd
from markdown import markdown

from config_manager import API_RETRY_DELAYS, DEFAULT_STYLES, MODEL_LIMITS

def retry_on_failure(func):
    @wraps(func)
    def wrapper(*args, **kwargs):
        attempts = len(API_RETRY_DELAYS) + 1
        for i in range(attempts):
            try:
                return func(*args, **kwargs)
            except Exception as e:
                if i < len(API_RETRY_DELAYS):
                    time.sleep(API_RETRY_DELAYS[i])
                else:
                    raise
    return wrapper

def calculate_max_input_tokens(model_name: str, requested_output_tokens: int) -> int:
    """
    Calcule le nombre de tokens d'entrée autorisés en soustrayant la sortie demandée
    et une marge de sécurité de la fenêtre de contexte totale du modèle.
    """
    model_info = MODEL_LIMITS.get(model_name)
    if not model_info:
        # Fallback sécuritaire pour des modèles inconnus
        return 4000
    
    total_context = model_info.get("context", 8000)
    # Marge de sécurité de 10% de la sortie, avec un minimum de 200 tokens
    safety_margin = max(200, int(0.1 * requested_output_tokens))
    
    allowed_input = total_context - requested_output_tokens - safety_margin
    return max(0, allowed_input)

@retry_on_failure
def call_openai(
    model_name: str, 
    prompt: str, 
    api_key: str,
    temperature: float = 0.7, 
    top_p: float = 0.9, 
    max_output_tokens: int = 1024,
    reasoning_effort: str = "medium",
    verbosity: str = "medium"
) -> str:
    from openai import OpenAI
    client = OpenAI(api_key=api_key)

    # Vérifier si c'est un modèle GPT-5 (qui supporte reasoning_effort et verbosity)
    is_gpt5 = "gpt-5" in model_name.lower()
    
    if is_gpt5:
        # Pour GPT-5, utiliser l'API Responses (recommandée) avec reasoning et text
        response = client.responses.create(
            model=model_name,
            input=prompt,
            max_output_tokens=max_output_tokens,     # Limite de sortie
            reasoning={"effort": reasoning_effort},  # Paramètre GPT-5
            text={"verbosity": verbosity}            # Paramètre GPT-5
        )
    else:
        # Pour les autres modèles, utiliser l'API Chat completions classique
        response = client.chat.completions.create(
            model=model_name,
            messages=[{"role": "user", "content": prompt}],
            temperature=temperature,
            top_p=top_p,
            max_tokens=max_output_tokens
        )
    
    # Récupération de la réponse selon l'API utilisée
    if is_gpt5:
        # Pour GPT-5 (API Responses)
        # Le contenu textuel est dans response.output
        if hasattr(response, "output") and response.output:
            # response.output est une liste avec des éléments de type ResponseOutputMessage
            for item in response.output:
                if hasattr(item, "type") and item.type == "message":
                    # C'est un message, chercher le contenu textuel
                    if hasattr(item, "content") and item.content:
                        for content_item in item.content:
                            if hasattr(content_item, "type") and content_item.type == "output_text":
                                if hasattr(content_item, "text"):
                                    return content_item.text or ""
        
        # Fallback: retourner la représentation string
        return str(response)
    else:
        # Pour les autres modèles (API Chat completions)
        if hasattr(response, "choices") and response.choices and hasattr(response.choices[0], "message"):
            return response.choices[0].message.content or ""
        elif hasattr(response, "choices") and response.choices and hasattr(response.choices[0], "text"):
            return response.choices[0].text or ""
    
    # Fallback pour le débuggage
    return str(response)

@retry_on_failure
def call_anthropic(
    model_name: str, 
    prompt: str, 
    api_key: str,
    temperature: float = 0.7, 
    top_p: float = 0.9, 
    max_output_tokens: int = 1024
) -> str:
    import anthropic
    client = anthropic.Anthropic(api_key=api_key)
    msg = client.messages.create(
        model=model_name,
        max_tokens=max_output_tokens, # S'assurer que ce paramètre est bien nommé 'max_tokens'
        temperature=temperature,
        top_p=top_p,
        messages=[{"role":"user","content":prompt}]
    )
    
    parts = []
    for blk in getattr(msg, "content", []) or []:
        if isinstance(blk, dict) and blk.get("type") == "text":
            parts.append(blk.get("text",""))
        else:
            try:
                parts.append(getattr(blk, "text", "") or "")
            except Exception:
                pass
    return "\n".join(parts).strip() or str(msg)

def parse_docx_plan(docx_path: str) -> List[Dict[str, Any]]:
    doc = Document(docx_path)
    heading_map = {"Heading 1":1,"Heading 2":2,"Heading 3":3,"Titre 1":1,"Titre 2":2,"Titre 3":3}
    counters = {1:0,2:0,3:0}
    items = []
    for p in doc.paragraphs:
        level = heading_map.get(getattr(getattr(p, "style", None), "name", ""))
        if level:
            for lv in [3,2,1]:
                if lv > level: counters[lv] = 0
            counters[level] += 1
            # Construire le code de section en incluant tous les niveaux
            # jusqu'au niveau courant (même si un niveau parent vaut 0)
            parts = [str(counters[lv]) for lv in range(1, level + 1)]
            code = ".".join(parts)
            items.append({"code": code, "title": p.text.strip(), "level": level})
    return items

def _hex_to_rgb(hex_str: str) -> RGBColor:
    """Convertit une chaîne hex (ex. '1F3864') en RGBColor."""
    h = hex_str.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def _apply_doc_styles(doc, styles: dict):
    """Configure les marges, le format de page et les styles de base du document."""
    section = doc.sections[0]
    # Marges
    section.top_margin = Cm(styles.get("margin_top", DEFAULT_STYLES["margin_top"]))
    section.bottom_margin = Cm(styles.get("margin_bottom", DEFAULT_STYLES["margin_bottom"]))
    section.left_margin = Cm(styles.get("margin_left", DEFAULT_STYLES["margin_left"]))
    section.right_margin = Cm(styles.get("margin_right", DEFAULT_STYLES["margin_right"]))
    # Taille de page (A4 par défaut)
    section.page_width = Cm(styles.get("page_width", DEFAULT_STYLES.get("page_width", 21.0)))
    section.page_height = Cm(styles.get("page_height", DEFAULT_STYLES.get("page_height", 29.7)))


def _setup_heading_style(doc, style_name: str, font_family: str, size: int,
                         color_hex: str, bold: bool, space_before: int, space_after: int):
    """Configure un style de titre Word natif (Heading 1/2/3)."""
    try:
        style = doc.styles[style_name]
    except KeyError:
        return
    font = style.font
    font.name = font_family
    font.size = Pt(size)
    font.bold = bold
    font.color.rgb = _hex_to_rgb(color_hex)
    # Forcer la police pour les caractères Est-Asie / complexes
    rpr = style.element.find(qn("w:rPr"))
    if rpr is None:
        rpr = style.element.makeelement(qn("w:rPr"), {})
        style.element.append(rpr)
    for tag in [qn("w:rFonts")]:
        existing = rpr.find(tag)
        if existing is not None:
            existing.set(qn("w:ascii"), font_family)
            existing.set(qn("w:hAnsi"), font_family)
            existing.set(qn("w:cs"), font_family)
        else:
            rfonts = rpr.makeelement(qn("w:rFonts"), {
                qn("w:ascii"): font_family,
                qn("w:hAnsi"): font_family,
                qn("w:cs"): font_family,
            })
            rpr.append(rfonts)
    pf = style.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)


def _setup_normal_style(doc, font_family: str, size: int, line_spacing: float,
                        space_after: int, first_line_indent: float):
    """Configure le style Normal (corps de texte)."""
    style = doc.styles["Normal"]
    font = style.font
    font.name = font_family
    font.size = Pt(size)
    font.color.rgb = RGBColor(0x26, 0x27, 0x30)  # gris très foncé, meilleur que noir pur
    pf = style.paragraph_format
    pf.line_spacing = line_spacing
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(0)
    if first_line_indent > 0:
        pf.first_line_indent = Cm(first_line_indent)
    # Forcer la police via XML pour compatibilité complète
    rpr = style.element.find(qn("w:rPr"))
    if rpr is None:
        rpr = style.element.makeelement(qn("w:rPr"), {})
        style.element.append(rpr)
    existing = rpr.find(qn("w:rFonts"))
    attrs = {
        qn("w:ascii"): font_family,
        qn("w:hAnsi"): font_family,
        qn("w:cs"): font_family,
        qn("w:eastAsia"): font_family,
    }
    if existing is not None:
        for k, v in attrs.items():
            existing.set(k, v)
    else:
        rfonts = rpr.makeelement(qn("w:rFonts"), attrs)
        rpr.append(rfonts)


def _setup_list_styles(doc, font_family: str, size: int, line_spacing: float):
    """Configure les styles de liste pour qu'ils héritent de la police/taille."""
    for sname in ("List Bullet", "List Number"):
        try:
            style = doc.styles[sname]
        except KeyError:
            continue
        style.font.name = font_family
        style.font.size = Pt(size)
        pf = style.paragraph_format
        pf.line_spacing = line_spacing
        pf.space_after = Pt(2)
        pf.space_before = Pt(1)


def _parse_inline_markdown(paragraph, text: str, font_family: str, size: int,
                           color: RGBColor = None):
    """Parse le Markdown inline (gras, italique) et ajoute des runs au paragraphe.

    Supporte : **gras**, *italique*, ***gras+italique***, `code inline`.
    """
    # Pattern pour capturer les segments formatés
    pattern = re.compile(
        r'(\*\*\*(.+?)\*\*\*)'   # ***bold+italic***
        r'|(\*\*(.+?)\*\*)'       # **bold**
        r'|(\*(.+?)\*)'           # *italic*
        r'|(`(.+?)`)'             # `code`
    )
    last_end = 0
    for m in pattern.finditer(text):
        # Texte brut avant le match
        if m.start() > last_end:
            _add_run(paragraph, text[last_end:m.start()], font_family, size, color=color)
        if m.group(2):    # ***bold+italic***
            _add_run(paragraph, m.group(2), font_family, size, bold=True, italic=True, color=color)
        elif m.group(4):  # **bold**
            _add_run(paragraph, m.group(4), font_family, size, bold=True, color=color)
        elif m.group(6):  # *italic*
            _add_run(paragraph, m.group(6), font_family, size, italic=True, color=color)
        elif m.group(8):  # `code`
            _add_run(paragraph, m.group(8), "Courier New", max(size - 1, 8), color=RGBColor(0x80, 0x30, 0x30))
        last_end = m.end()
    # Texte restant
    if last_end < len(text):
        _add_run(paragraph, text[last_end:], font_family, size, color=color)


def _add_run(paragraph, text: str, font_family: str, size: int,
             bold: bool = False, italic: bool = False,
             color: RGBColor = None):
    """Ajoute un run avec mise en forme complète."""
    r = paragraph.add_run(text)
    r.font.name = font_family
    r.font.size = Pt(size)
    if bold:
        r.font.bold = True
    if italic:
        r.font.italic = True
    if color:
        r.font.color.rgb = color
    return r


def _add_styled_paragraph(doc, text: str, font_family: str, size: int,
                          line_spacing: float = None, color: RGBColor = None):
    """Ajoute un paragraphe avec parsing Markdown inline."""
    p = doc.add_paragraph()
    _parse_inline_markdown(p, text, font_family, size, color=color)
    if line_spacing:
        p.paragraph_format.line_spacing = line_spacing
    return p


def generate_styled_docx(markdown_text: str, output_path: str, styles: Dict[str, Any]) -> None:
    """Génère un document DOCX de haute qualité à partir de Markdown.

    Améliorations par rapport à la version basique :
    - Styles Word natifs (Heading 1-3, Normal, List Bullet, List Number)
    - Parsing inline Markdown (**gras**, *italique*, `code`)
    - Interligne, espacement paragraphe, couleurs titres
    - Format de page A4 avec marges configurables
    - Listes numérotées (1. 2. 3.)
    """
    # Lecture de la configuration de style
    s = lambda key: styles.get(key, DEFAULT_STYLES.get(key))
    font_family = s("font_family") or "Calibri"
    body_size = int(s("font_size_body") or 11)
    h1_size = int(s("font_size_h1") or 18)
    h2_size = int(s("font_size_h2") or 14)
    h3_size = int(s("font_size_h3") or 12)
    line_spacing = float(s("line_spacing") or 1.15)
    space_after = int(s("space_after_paragraph") or 6)
    heading_bold = bool(s("heading_bold") if s("heading_bold") is not None else True)
    color_h1 = s("heading_color_h1") or "1F3864"
    color_h2 = s("heading_color_h2") or "2E5090"
    color_h3 = s("heading_color_h3") or "404040"
    first_indent = float(s("first_line_indent") or 0)

    # Création du document et configuration globale
    doc = Document()
    _apply_doc_styles(doc, styles)

    # Configuration des styles natifs Word
    _setup_normal_style(doc, font_family, body_size, line_spacing, space_after, first_indent)
    _setup_heading_style(doc, "Heading 1", font_family, h1_size, color_h1, heading_bold,
                         int(s("space_before_h1") or 24), int(s("space_after_h1") or 12))
    _setup_heading_style(doc, "Heading 2", font_family, h2_size, color_h2, heading_bold,
                         int(s("space_before_h2") or 18), int(s("space_after_h2") or 8))
    _setup_heading_style(doc, "Heading 3", font_family, h3_size, color_h3, heading_bold,
                         int(s("space_before_h3") or 12), int(s("space_after_h3") or 6))
    _setup_list_styles(doc, font_family, body_size, line_spacing)

    # Regex pour listes numérotées (ex: "1. Texte")
    re_numbered = re.compile(r"^(\d+)\.\s+(.*)")

    # Parsing ligne par ligne
    blank_count = 0
    for line in markdown_text.splitlines():
        stripped = line.strip()

        # Lignes vides : une seule ligne vide = espacement, multiples = 1 seul saut
        if not stripped:
            blank_count += 1
            if blank_count <= 1:
                p = doc.add_paragraph("")
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(2)
            continue
        blank_count = 0

        # --- Titres (styles natifs Word) ---
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
            continue
        if stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
            continue
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
            continue

        # --- Listes à puces ---
        if stripped.startswith(("- ", "* ")):
            p = doc.add_paragraph(style="List Bullet")
            _parse_inline_markdown(p, stripped[2:], font_family, body_size)
            continue

        # --- Listes numérotées ---
        m_num = re_numbered.match(stripped)
        if m_num:
            p = doc.add_paragraph(style="List Number")
            _parse_inline_markdown(p, m_num.group(2), font_family, body_size)
            continue

        # --- Séparateur horizontal ---
        if stripped in ("---", "***", "___"):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            # Ajout d'une bordure basse via XML
            pPr = p._p.get_or_add_pPr()
            pBdr = pPr.makeelement(qn("w:pBdr"), {})
            bottom = pBdr.makeelement(qn("w:bottom"), {
                qn("w:val"): "single",
                qn("w:sz"): "6",
                qn("w:space"): "1",
                qn("w:color"): "CCCCCC",
            })
            pBdr.append(bottom)
            pPr.append(pBdr)
            continue

        # --- Paragraphe normal avec Markdown inline ---
        _add_styled_paragraph(doc, stripped, font_family, body_size, line_spacing)

    doc.save(output_path)

def extract_used_references_apa(text_md: str) -> List[str]:
    pats = re.findall(r"\(([^,]+),\s*(\d{4})\)", text_md)
    return sorted({f"{a}, {y}" for a,y in pats})

def truncate_to_tokens(text: str, max_tokens: int, model: str = "gpt-4") -> str:
    """Tronque le texte selon le nombre max de tokens.
    Utilise tiktoken si disponible, sinon heuristique ~4 caractères = 1 token."""
    try:
        import tiktoken
        enc = tiktoken.get_encoding("cl100k_base")
        tokens = enc.encode(text)
        if len(tokens) <= max_tokens:
            return text
        return enc.decode(tokens[:max_tokens])
    except Exception:
        # Fallback heuristique : ~4 caractères par token
        max_chars = max(256, int(max_tokens * 4))
        return text[:max_chars] if len(text) > max_chars else text

def _generate_filename(base_name: str, mode: str) -> str:
    """Génère un nom de fichier avec timestamp."""
    timestamp = datetime.now().strftime("%Y%m%d-%H%M%S")
    # Nettoie le base_name pour éviter les caractères problématiques
    slug = "".join(c if c.isalnum() or c in "-_" else "-" for c in base_name)[:80]
    return f"{timestamp}_{mode}_{slug or 'sortie'}"

def export_markdown(text_md: str, base_name: str, mode: str, export_dir: str = "output") -> str:
    """Exporte le texte markdown dans un fichier .md."""
    outdir = Path(export_dir)
    outdir.mkdir(parents=True, exist_ok=True)
    filename = _generate_filename(base_name, mode)
    path = outdir / f"{filename}.md"
    path.write_text(text_md, encoding="utf-8")
    return str(path)

def export_docx(text_md: str, base_name: str, mode: str, export_dir: str = "output", styles: dict = None) -> str:
    """Exporte le texte markdown dans un fichier .docx stylé."""
    outdir = Path(export_dir)
    outdir.mkdir(parents=True, exist_ok=True)
    filename = _generate_filename(base_name, mode)
    path = outdir / f"{filename}.docx"
    generate_styled_docx(text_md, str(path), styles or {})
    return str(path)

def generate_bibliography(used: List[str], excel_path: str) -> str:
    import pandas as pd
    try:
        df = pd.read_excel(excel_path, sheet_name="Bibliographie")
    except Exception:
        df = pd.read_csv(excel_path)
    cols = {c.lower(): c for c in df.columns}
    col_full = cols.get("référence apa complète") or cols.get("reference apa complete") or cols.get("apa_full")
    col_short = cols.get("référence courte") or cols.get("reference courte") or cols.get("apa_short")
    if not col_full:
        return "\n".join(f"- {ref}" for ref in used)
    lines = []
    shorts = set(used)
    if col_short and col_short in df.columns:
        m = df[df[col_short].astype(str).str.strip().isin(shorts)]
        for _, r in m.iterrows(): lines.append(f"- {r[col_full]}")
    else:
        for ref in used:
            year = ref.split(",")[-1].strip()
            author = ref.rsplit(",",1)[0].strip()
            mask = df[col_full].astype(str).str.contains(author) & df[col_full].astype(str).str.contains(year)
            if mask.any(): lines.append(f"- {df[mask].iloc[0][col_full]}")
            else: lines.append(f"- {ref}")
    return "\n".join(lines)
