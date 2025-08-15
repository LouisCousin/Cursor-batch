#!/usr/bin/env python3
"""
Module d'utilitaires d'export pour les prompts et résultats.
Gère l'export des prompts de génération au format Word.
"""

import os
from datetime import datetime
from typing import List, Dict, Any, Optional
from pathlib import Path
import logging

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    Document = None
    logging.warning("python-docx non disponible. Fonctionnalités d'export Word désactivées.")


class PromptExporter:
    """
    Gestionnaire d'export des prompts au format Word.
    """
    
    def __init__(self):
        """Initialise l'exporteur."""
        self.docx_available = DOCX_AVAILABLE
        
    def create_prompts_document(self, prompts_data: List[Dict[str, Any]], 
                              export_path: str, 
                              title: str = "Prompts de Génération",
                              include_corpus: bool = False) -> Optional[str]:
        """
        Crée un document Word contenant tous les prompts.
        
        Args:
            prompts_data: Liste des données de prompts à exporter
            export_path: Chemin de destination du fichier
            title: Titre du document
            include_corpus: Inclure ou non le contenu du corpus
            
        Returns:
            Chemin du fichier créé ou None en cas d'erreur
        """
        if not self.docx_available:
            logging.error("python-docx non disponible pour l'export Word")
            return None
        
        try:
            # Créer le document
            doc = Document()
            
            # Configurer les styles
            self._setup_document_styles(doc)
            
            # Ajouter le titre principal
            title_para = doc.add_heading(title, level=0)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Ajouter les métadonnées
            doc.add_paragraph(f"Généré le : {datetime.now().strftime('%d/%m/%Y à %H:%M')}")
            doc.add_paragraph(f"Nombre de prompts : {len(prompts_data)}")
            doc.add_paragraph(f"Contenu du corpus inclus : {'Oui' if include_corpus else 'Non'}")
            doc.add_paragraph("")  # Ligne vide
            
            # Ajouter une table des matières simplifiée
            doc.add_heading("Table des Matières", level=1)
            toc_para = doc.add_paragraph()
            for i, prompt_info in enumerate(prompts_data, 1):
                section_title = prompt_info.get('section_title', f'Section {i}')
                section_code = prompt_info.get('section_code', '')
                toc_line = f"{i}. {section_code} - {section_title}"
                toc_para.add_run(toc_line + "\n")
            
            doc.add_page_break()
            
            # Ajouter chaque prompt
            for i, prompt_info in enumerate(prompts_data, 1):
                self._add_prompt_section(doc, prompt_info, i, include_corpus)
                
                # Ajouter un saut de page sauf pour le dernier
                if i < len(prompts_data):
                    doc.add_page_break()
            
            # Ajouter un pied de document
            doc.add_paragraph("")
            footer_para = doc.add_paragraph("---")
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            footer_para = doc.add_paragraph(f"Document généré automatiquement par le Générateur d'Ouvrage Assisté par IA")
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Sauvegarder le document
            os.makedirs(os.path.dirname(export_path), exist_ok=True)
            doc.save(export_path)
            
            logging.info(f"Document de prompts exporté vers : {export_path}")
            return export_path
            
        except Exception as e:
            logging.error(f"Erreur lors de l'export des prompts : {e}")
            return None
    
    def _setup_document_styles(self, doc: Document):
        """Configure les styles du document."""
        try:
            # Style pour les titres de section
            styles = doc.styles
            
            # Style pour le code/prompt
            if 'Code' not in [style.name for style in styles]:
                code_style = styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)
                code_font = code_style.font
                code_font.name = 'Courier New'
                code_font.size = Pt(9)
                code_style.paragraph_format.left_indent = Inches(0.5)
                code_style.paragraph_format.space_after = Pt(6)
            
            # Style pour les métadonnées
            if 'Metadata' not in [style.name for style in styles]:
                meta_style = styles.add_style('Metadata', WD_STYLE_TYPE.PARAGRAPH)
                meta_font = meta_style.font
                meta_font.italic = True
                meta_font.size = Pt(10)
                meta_style.paragraph_format.space_after = Pt(3)
                
        except Exception as e:
            logging.warning(f"Impossible de configurer les styles personnalisés : {e}")
    
    def _add_prompt_section(self, doc: Document, prompt_info: Dict[str, Any], 
                           section_number: int, include_corpus: bool):
        """
        Ajoute une section de prompt au document.
        
        Args:
            doc: Document Word
            prompt_info: Informations sur le prompt
            section_number: Numéro de la section
            include_corpus: Inclure le corpus dans le prompt
        """
        # Titre de la section
        section_code = prompt_info.get('section_code', '')
        section_title = prompt_info.get('section_title', f'Section {section_number}')
        
        heading = doc.add_heading(f"{section_number}. {section_code} - {section_title}", level=1)
        
        # Métadonnées de la section
        if 'timestamp' in prompt_info:
            try:
                meta_para = doc.add_paragraph(f"Généré le : {prompt_info['timestamp']}")
                meta_para.style = 'Metadata'
            except:
                doc.add_paragraph(f"Généré le : {prompt_info['timestamp']}")
        
        if 'model_used' in prompt_info:
            try:
                meta_para = doc.add_paragraph(f"Modèle utilisé : {prompt_info['model_used']}")
                meta_para.style = 'Metadata'
            except:
                doc.add_paragraph(f"Modèle utilisé : {prompt_info['model_used']}")
        
        doc.add_paragraph("")  # Ligne vide
        
        # Sous-titre pour le prompt
        doc.add_heading("Prompt de Génération", level=2)
        
        # Contenu du prompt
        prompt_text = prompt_info.get('prompt', 'Prompt non disponible')
        
        if not include_corpus:
            # Nettoyer le prompt du contenu du corpus
            prompt_text = self._clean_prompt_from_corpus(prompt_text)
        
        # Ajouter le prompt avec style code
        try:
            prompt_para = doc.add_paragraph(prompt_text)
            prompt_para.style = 'Code'
        except:
            # Fallback si le style n'est pas disponible
            prompt_para = doc.add_paragraph(prompt_text)
            prompt_run = prompt_para.runs[0]
            prompt_run.font.name = 'Courier New'
            prompt_run.font.size = Pt(9)
        
        # Informations sur le corpus si disponible
        if 'corpus_info' in prompt_info:
            doc.add_paragraph("")
            doc.add_heading("Informations sur le Corpus", level=2)
            
            corpus_info = prompt_info['corpus_info']
            if isinstance(corpus_info, dict):
                doc.add_paragraph(f"Nombre d'entrées utilisées : {corpus_info.get('count', 'N/A')}")
                doc.add_paragraph(f"Score de pertinence moyen : {corpus_info.get('avg_score', 'N/A')}")
                if 'sources' in corpus_info:
                    doc.add_paragraph("Sources principales :")
                    for source in corpus_info['sources'][:5]:  # Limiter à 5 sources
                        doc.add_paragraph(f"• {source}", style='List Bullet')
        
        # Résultat de la génération (optionnel)
        if 'result_preview' in prompt_info and prompt_info['result_preview']:
            doc.add_paragraph("")
            doc.add_heading("Aperçu du Résultat", level=2)
            result_preview = prompt_info['result_preview']
            if len(result_preview) > 500:
                result_preview = result_preview[:500] + "... [Tronqué]"
            doc.add_paragraph(result_preview)
    
    def _clean_prompt_from_corpus(self, prompt_text: str) -> str:
        """
        Nettoie le prompt du contenu du corpus pour ne garder que la structure.
        
        Args:
            prompt_text: Texte du prompt original
            
        Returns:
            Prompt nettoyé
        """
        # Patterns courants à nettoyer
        patterns_to_clean = [
            "--- CORPUS DATA ---",
            "--- END CORPUS DATA ---",
            "DONNÉES DU CORPUS :",
            "FIN DES DONNÉES DU CORPUS",
        ]
        
        lines = prompt_text.split('\n')
        cleaned_lines = []
        in_corpus_section = False
        
        for line in lines:
            line_lower = line.lower().strip()
            
            # Détecter le début d'une section corpus
            if any(pattern.lower() in line_lower for pattern in patterns_to_clean):
                if "fin" in line_lower or "end" in line_lower:
                    in_corpus_section = False
                    cleaned_lines.append("[... Contenu du corpus retiré ...]")
                else:
                    in_corpus_section = True
                continue
            
            # Heuristique : si une ligne contient beaucoup de données structurées
            if in_corpus_section:
                continue
            
            # Détecter les blocs de données JSON/CSV
            if (line.strip().startswith('{') or 
                line.strip().startswith('[') or
                '","' in line or
                line.count('|') > 3):
                if not in_corpus_section:
                    cleaned_lines.append("[... Données du corpus ...]")
                    in_corpus_section = True
                continue
            
            # Si on n'est pas dans une section corpus, garder la ligne
            if not in_corpus_section:
                cleaned_lines.append(line)
            else:
                # Sortir du mode corpus si on trouve une ligne normale
                if line.strip() and not any(char in line for char in ['{', '}', '[', ']', '","']):
                    in_corpus_section = False
                    cleaned_lines.append(line)
        
        return '\n'.join(cleaned_lines)
    
    def export_generation_results_prompts(self, generation_results: Dict[str, Any],
                                        export_dir: str = "output",
                                        filename: str = None) -> Optional[str]:
        """
        Exporte les prompts des résultats de génération.
        
        Args:
            generation_results: Résultats de génération de la session Streamlit
            export_dir: Dossier de destination
            filename: Nom du fichier (optionnel)
            
        Returns:
            Chemin du fichier créé ou None
        """
        if not generation_results:
            logging.warning("Aucun résultat de génération à exporter")
            return None
        
        # Préparer les données des prompts
        prompts_data = []
        
        for section_key, result in generation_results.items():
            prompt_info = {
                'section_code': result.get('section_code', ''),
                'section_title': result.get('section_title', ''),
                'prompt': result.get('original_prompt', 'Prompt non disponible'),
                'timestamp': result.get('timestamp', ''),
                'model_used': result.get('model_used', 'Non spécifié'),
                'result_preview': result.get('brouillon', '')[:500] if result.get('brouillon') else None
            }
            prompts_data.append(prompt_info)
        
        # Générer le nom de fichier si non fourni
        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d-%H%M%S")
            filename = f"{timestamp}_prompts_generation.docx"
        
        # Chemin complet
        export_path = os.path.join(export_dir, filename)
        
        # Créer le document
        return self.create_prompts_document(
            prompts_data, 
            export_path, 
            title="Prompts de Génération - Brouillons (IA 1)",
            include_corpus=False
        )
    
    def export_batch_prompts(self, batch_data: List[Dict[str, Any]],
                           export_dir: str = "output",
                           batch_id: str = None) -> Optional[str]:
        """
        Exporte les prompts d'un batch de traitement.
        
        Args:
            batch_data: Données du batch à exporter
            export_dir: Dossier de destination
            batch_id: ID du batch (pour le nom de fichier)
            
        Returns:
            Chemin du fichier créé ou None
        """
        if not batch_data:
            return None
        
        # Générer le nom de fichier
        timestamp = datetime.now().strftime("%Y%m%d-%H%M%S")
        if batch_id:
            filename = f"{timestamp}_batch_{batch_id}_prompts.docx"
        else:
            filename = f"{timestamp}_batch_prompts.docx"
        
        export_path = os.path.join(export_dir, filename)
        
        return self.create_prompts_document(
            batch_data,
            export_path,
            title=f"Prompts du Batch {batch_id}" if batch_id else "Prompts de Batch",
            include_corpus=False
        )


def create_export_summary_document(results: Dict[str, Any], 
                                 export_path: str,
                                 process_info: Dict[str, Any] = None) -> Optional[str]:
    """
    Crée un document de synthèse des résultats d'export.
    
    Args:
        results: Résultats de génération
        export_path: Chemin de destination
        process_info: Informations sur le processus (optionnel)
        
    Returns:
        Chemin du fichier créé ou None
    """
    if not DOCX_AVAILABLE:
        return None
    
    try:
        doc = Document()
        
        # Titre
        title = doc.add_heading("Synthèse de Génération d'Ouvrage", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Métadonnées
        doc.add_paragraph(f"Généré le : {datetime.now().strftime('%d/%m/%Y à %H:%M')}")
        if process_info:
            doc.add_paragraph(f"Processus : {process_info.get('process_id', 'N/A')}")
            doc.add_paragraph(f"Type : {process_info.get('type', 'N/A')}")
        
        doc.add_paragraph("")
        
        # Statistiques
        doc.add_heading("Statistiques", level=1)
        doc.add_paragraph(f"Nombre total de sections : {len(results)}")
        
        completed_sections = len([r for r in results.values() if 'brouillon' in r])
        finalized_sections = len([r for r in results.values() if 'finale' in r])
        
        doc.add_paragraph(f"Sections avec brouillon : {completed_sections}")
        doc.add_paragraph(f"Sections finalisées : {finalized_sections}")
        doc.add_paragraph(f"Taux de complétion : {(completed_sections/len(results)*100):.1f}%")
        
        # Liste des sections
        doc.add_heading("Sections Traitées", level=1)
        
        for section_key, result in results.items():
            doc.add_heading(f"{result.get('section_code', '')} - {result.get('section_title', '')}", level=2)
            
            status = "✅ Finalisée" if 'finale' in result else ("📝 Brouillon" if 'brouillon' in result else "❌ Non traitée")
            doc.add_paragraph(f"Statut : {status}")
            
            if 'timestamp' in result:
                doc.add_paragraph(f"Traité le : {result['timestamp']}")
            
            if 'md_path' in result:
                doc.add_paragraph(f"Fichier MD : {os.path.basename(result['md_path'])}")
            
            if 'docx_path' in result:
                doc.add_paragraph(f"Fichier DOCX : {os.path.basename(result['docx_path'])}")
            
            doc.add_paragraph("")
        
        # Sauvegarder
        doc.save(export_path)
        return export_path
        
    except Exception as e:
        logging.error(f"Erreur lors de la création du document de synthèse : {e}")
        return None
